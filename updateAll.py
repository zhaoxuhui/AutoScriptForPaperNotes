from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor  # 设置字体颜色
import docx
import time
import os
import sys
import win32com.client
import re
import datetime
import calendar
from matplotlib import pyplot as plt


def findAllFiles(root_dir, filter):
    """
    在指定目录查找指定类型文件

    :param root_dir: 查找目录
    :param filter: 文件类型
    :return: 路径、名称、文件全路径

    """

    print("Finding files ends with \'" + filter + "\' ...")
    separator = os.path.sep
    paths = []
    names = []
    files = []
    for parent, dirname, filenames in os.walk(root_dir):
        for filename in filenames:
            if filename.endswith(filter):
                paths.append(parent + separator)
                names.append(filename)
    for i in range(paths.__len__()):
        files.append(paths[i] + names[i])
    print(names.__len__().__str__() + " files have been found.")
    return paths, names, files


def formatDate(str_date):
    """
    用于格式化字符串
    :param str_date: 未格式化字符串
    :return: xxxx.xx.xx格式日期
    """
    split_parts = []
    # 2020.7.13
    if str_date.__contains__('.'):
        split_parts = str_date.split('.')
    # 2020-7-13
    elif str_date.__contains__('-'):
        split_parts = str_date.split('-')
    # 2020年7月13日
    elif str_date.__contains__('年'):
        tmp_year = str_date.split('年')[0]
        tmp_month = str_date.split('年')[1].split('月')[0]
        tmp_day = str_date.split('年')[1].split('月')[1].split('日')[0]
        split_parts.append(tmp_year)
        split_parts.append(tmp_month)
        split_parts.append(tmp_day)

    year = split_parts[0]
    month = split_parts[1].zfill(2)
    day = split_parts[2].zfill(2)
    return year + '.' + month + '.' + day


def parseDateToNumber(str_date):
    split_parts = str_date.split(".")
    num_year = int(split_parts[0])
    num_month = int(split_parts[1])
    num_day = int(split_parts[2])
    return num_year, num_month, num_day


def readNoteFiles(file_list, name_list):
    """
    根据找到的笔记文件读取内容
    :param file_list: 笔记文件路径列表
    :param name_list: 文章的真实名称列表
    :return: 解析的一系列内容
    """
    title_list = []
    author_list = []
    keyword_list = []
    publisher_list = []
    publish_time_list = []
    reading_time_list = []
    comment_list = []
    expression_list = []
    importantWord_list = []
    sorted_name_list = []

    for i in range(len(file_list)):
        f_in = open(file_list[i], 'r', encoding='UTF-8')
        f_in.readline()
        title_list.append(f_in.readline().strip())
        f_in.readline()
        author_list.append(f_in.readline().strip())
        f_in.readline()
        keyword_list.append(f_in.readline().strip())
        f_in.readline()
        publisher_list.append(f_in.readline().strip())
        f_in.readline()
        publish_time_list.append(f_in.readline().strip())
        f_in.readline()
        reading_time_list.append(formatDate(f_in.readline().strip()))
        f_in.readline()
        str_comment = f_in.readline()
        tmp_str = f_in.readline()
        while not tmp_str.__contains__("Great Expressions:"):
            str_comment += tmp_str
            tmp_str = f_in.readline()
        comment_list.append(str_comment.strip())

        str_expression = f_in.readline()
        tmp_str = f_in.readline()
        while not tmp_str.__contains__("Important Words:"):
            str_expression += tmp_str
            tmp_str = f_in.readline()
        expression_list.append(str_expression.strip())

        str_words = f_in.readline()
        tmp_str = f_in.readline()
        while tmp_str is not "":
            str_words += tmp_str
            tmp_str = f_in.readline()
        importantWord_list.append(str_words.strip())

        f_in.close()

    data = [(reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name) for
            reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name in
            zip(reading_time_list, title_list, author_list, keyword_list, publisher_list, publish_time_list,
                comment_list, expression_list, importantWord_list, name_list)]
    data.sort()  # 按照阅读日期进行排序

    title_list = [title for
                  reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name in
                  data]
    author_list = [author for
                   reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name
                   in
                   data]
    keyword_list = [keyword for
                    reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name
                    in
                    data]
    publisher_list = [publisher for
                      reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name
                      in
                      data]
    publish_time_list = [publish_time for
                         reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name
                         in
                         data]
    reading_time_list = [reading_time for
                         reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name
                         in
                         data]
    comment_list = [comment for
                    reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name
                    in
                    data]
    expression_list = [expression for
                       reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name
                       in
                       data]
    importantWord_list = [word for
                          reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name
                          in
                          data]
    sorted_name_list = [file_name for
                        reading_time, title, author, keyword, publisher, publish_time, comment, expression, word, file_name
                        in
                        data]

    return title_list, author_list, \
           keyword_list, publisher_list, \
           publish_time_list, reading_time_list, \
           comment_list, expression_list, \
           importantWord_list, sorted_name_list


def writeItemStyle1(document,
                    str_title,
                    str_author,
                    str_keyword,
                    str_publisher,
                    str_time,
                    str_readingDate,
                    str_comments):
    """
    包含文章标题、作者、关键词、出版社、出版时间、阅读日期、评论的格式内容
    """
    heading_title = document.add_heading(str_title, level=2)
    heading_author = document.add_heading("Authors", level=3)
    para_author = document.add_paragraph(str_author)
    if str_keyword is not '':
        heading_keyword = document.add_heading("Keywords", level=3)
        para_keywords = document.add_paragraph(str_keyword)
    heading_publisher = document.add_heading("Publisher & Time", level=3)
    str_content = str_publisher + ", " + str_time + """ <a href="https://scholar.google.com.sg/scholar?&q=""" + str_title + """"><Google Scholar></a>"""
    add_text_link(document, str_content)
    if str_comments is not '' or str_readingDate is not '':
        heading_readingTime = document.add_heading("Reading Date & Comments", level=3)
        para_readingTime = document.add_paragraph(str_readingDate + "\n" + str_comments)


def writeItemStyle2(document,
                    str_title,
                    str_author,
                    str_keyword,
                    str_publisher,
                    str_time,
                    str_readingDate,
                    str_comments,
                    str_expression,
                    str_words):
    """
    包含文章标题、作者、关键词、出版社、出版时间、阅读日期、评论、好的表达、重要单词的格式内容
    """
    heading_title = document.add_heading(str_title, level=2)
    heading_author = document.add_heading("Authors", level=3)
    para_author = document.add_paragraph(str_author)
    if str_keyword is not '':
        heading_keyword = document.add_heading("Keywords", level=3)
        para_keywords = document.add_paragraph(str_keyword)
    heading_publisher = document.add_heading("Publisher & Time", level=3)
    str_content = str_publisher + ", " + str_time + """ <a href="https://scholar.google.com.sg/scholar?&q=""" + str_title + """"><Google Scholar></a>"""
    add_text_link(document, str_content)
    if str_comments is not '' or str_readingDate is not '':
        heading_readingTime = document.add_heading("Reading Date & Comments", level=3)
        para_readingTime = document.add_paragraph(str_readingDate + "\n" + str_comments)
    if str_expression is not '':
        heading_expression = document.add_heading("Great Expressions", level=3)
        para_expression = document.add_paragraph(str_expression)
    if str_words is not '':
        heading_important = document.add_heading("Important Words", level=3)
        para_important = document.add_paragraph(str_words)


def writeHeading(document, str_heading):
    heading = document.add_heading()
    heading.alignment = 1
    heading_run = heading.add_run(str_heading)
    heading_run.font.size = Pt(20)
    heading_run.font.name = u'等线'
    heading_run.element.rPr.rFonts.set(qn('w:eastAsia'), u'等线')


def writeStartHeading(document, str_heading, total_num, str_start, str_end):
    """
    用于输出文档的主标题
    :param document: 文档对象
    :param str_heading: 标题内容
    :param total_num: 包含的文章个数
    :param str_start: 起始时间
    :param str_end: 结束时间
    :return: 无
    """
    writeHeading(document, str_heading)

    t = time.strftime('%Y.%m.%d', time.localtime(time.time()))
    str_content = "Auto-generated by Zhao Xuhui on " + t
    paragraph = document.add_paragraph(str_content)
    paragraph.alignment = 2
    str_content = "A total number of " + str(total_num) + " papers. "
    str_content += "(" + str_start + " - " + str_end + ")"
    paragraph = document.add_paragraph(str_content)


def outputExpressionAndWord(reading_time_list, name_list, expression_list, importantWord_list):
    """
    用于writing文件夹的输出
    """
    save_path_exp = "F:\\zhaoxuhui.github.io\\writing\\expressions\\"
    save_path_word = "F:\\zhaoxuhui.github.io\\writing\\words\\"
    for i in range(len(name_list)):
        if expression_list[i] is not "":
            path = save_path_exp + reading_time_list[i] + "_" + name_list[i][:50] + ".exp"
            # 如果文件不存在，新建，否则就跳过
            if not os.path.exists(path):
                fout = open(path, "w", encoding='UTF-8')
                fout.write(expression_list[i])
                fout.close()
            # 如果文件存在，检查内容是否一致，如果不一致删掉现在的，重新写入
            else:
                test_in = open(path, "r", encoding='UTF-8')
                test_lines = test_in.readlines()
                test_in.close()
                str_lines = ""
                for j in range(len(test_lines)):
                    str_lines += test_lines[j]
                if str_lines != expression_list[i]:
                    os.remove(path)
                    fout = open(path, "w", encoding='UTF-8')
                    fout.write(expression_list[i])
                    fout.close()

        if importantWord_list[i] is not "":
            path = save_path_word + reading_time_list[i] + "_" + name_list[i][:50] + ".wd"
            # 如果文件不存在，新建，否则就跳过
            if not os.path.exists(path):
                fout = open(path, "w", encoding='UTF-8')
                fout.write(importantWord_list[i])
                fout.close()
            # 如果文件存在，检查内容是否一致，如果不一致删掉现在的，重新写入
            else:
                test_in = open(path, "r", encoding='UTF-8')
                test_lines = test_in.readlines()
                test_in.close()
                str_lines = ""
                for j in range(len(test_lines)):
                    str_lines += test_lines[j]
                if str_lines != importantWord_list[i]:
                    os.remove(path)
                    fout = open(path, "w", encoding='UTF-8')
                    fout.write(importantWord_list[i])
                    fout.close()


def add_hyperlink(paragraph, url, text, color):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


# 判断字段是否为链接
def is_text_link(text):
    for i in ['http', '://', 'www.', '.com', '.org', '.cn', '.xyz', '.htm']:
        if i in text:
            return True
        else:
            return False


# 对段落中的链接加上超链接
def add_text_link(document, text):
    paragraph = document.add_paragraph()
    # 根据<a>标签拆分文本内容
    text = re.split(r'<a href="|">|</a>', text)
    keyword = None
    for i in range(len(text)):
        # 对非链接和非关键词的内容，通过run直接加入段落中
        if not is_text_link(text[i]):
            if text[i] != keyword:
                paragraph.add_run(text[i])
        # 对链接和关键词，使用add_hyperlink插入超链接
        elif i + 1 < len(text):
            url = text[i]
            keyword = text[i + 1]
            add_hyperlink(paragraph, url, keyword, '#ED7D31')


def strCvtToTimeStamp(timeStr):
    """
    将时间转换为时间戳
    """
    timeArray = time.strptime(timeStr, "%Y.%m.%d")
    timeStamp = int(time.mktime(timeArray))
    return timeStamp


def numCvtToTimeStamp(num_year, num_month, num_day):
    tmp_str = str(num_year) + "." + str(num_month) + "." + str(num_day)
    return strCvtToTimeStamp(tmp_str)


def numCvtToTimeStamp2(num_array):
    num_year = num_array[0]
    num_month = num_array[1]
    num_day = num_array[2]
    tmp_str = str(num_year) + "." + str(num_month) + "." + str(num_day)
    return strCvtToTimeStamp(tmp_str)


def timeStampCvtToStr(timeStamp):
    """
    时间戳转换为字符串
    """
    timeArray = time.localtime(timeStamp)
    strTime = time.strftime("%Y.%m.%d", timeArray)
    return strTime


def getIndexRange(dates, start_str, end_str):
    """
    获取指定日期所对应的列表索引
    """
    start_timestamp = strCvtToTimeStamp(start_str)
    end_timestamp = strCvtToTimeStamp(end_str)
    dates_timestamp = []
    for i in range(len(dates)):
        dates_timestamp.append(strCvtToTimeStamp(dates[i]))

    start_index = 0
    end_index = len(dates)
    if end_timestamp < start_timestamp:
        print("error time duration,return all")
    elif start_timestamp == end_timestamp:
        for i in range(len(dates_timestamp)):
            if dates_timestamp[i] == start_timestamp:
                start_index = end_index = i
                break
        if start_index == end_index:
            print("one particular day")
        else:
            print("one day - no matching day,return all")
    else:
        if end_timestamp < dates_timestamp[0]:
            print('too early,return all')
        elif start_timestamp > dates_timestamp[-1]:
            print("too late,return all")
        else:
            for i in range(len(dates_timestamp)):
                if dates_timestamp[i] >= start_timestamp:
                    start_index = i
                    break

            for i in range(len(dates_timestamp)):
                if dates_timestamp[i] == end_timestamp:
                    end_index = i + 1
                    break
                if dates_timestamp[i] > end_timestamp:
                    end_index = i
                    break
    return start_index, end_index


if __name__ == '__main__':
    select_flag = 0

    # 获取脚本所在的当前路径
    cur_dir = os.path.abspath(sys.argv[0])[:os.path.abspath(sys.argv[0]).rfind(os.sep)]
    # 如果没有传入任何参数，当前路径就作为搜素路径，否则按照输入的路径搜索
    # Case 1:没有传入任何参数，按照默认进行(默认路径，全部文件)
    if len(sys.argv) == 1:
        search_dir1 = cur_dir + "\\Closely Relevant"
        search_dir2 = cur_dir + "\\Generally Relevant"
    # Case 2:传入了1个参数，即搜索路径(自定义路径，全部文件)
    elif len(sys.argv) == 2:
        search_dir1 = sys.argv[1] + "\\Closely Relevant"
        search_dir2 = sys.argv[1] + "\\Generally Relevant"
    # Case 3:传入了2个参数，即起始时间和终止时间(默认路径，部分文件)
    elif len(sys.argv) == 3:
        search_dir1 = cur_dir + "\\Closely Relevant"
        search_dir2 = cur_dir + "\\Generally Relevant"
        print(search_dir1, search_dir2)
        start_date = formatDate(sys.argv[1])
        end_date = formatDate(sys.argv[2])
        select_flag = 1
    # Case 4:传入了3个参数，分别为搜索路径、起始时间、终止时间(自定义路径，部分文件)
    elif len(sys.argv) == 4:
        search_dir1 = sys.argv[1] + "\\Closely Relevant"
        search_dir2 = sys.argv[1] + "\\Generally Relevant"
        start_date = formatDate(sys.argv[2])
        end_date = formatDate(sys.argv[3])
        select_flag = 1

    # 构建文档对象
    document = Document('Template.docx')
    # 文档格式设置
    document.styles['Normal'].font.name = u'等线'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线')

    # 笔记文件的后缀名
    filter_type = ".note"

    # ----------------------------------------------------------------------------
    # Closely Relevant Papers
    # 开始查找所有笔记
    paths, names, files = findAllFiles(search_dir1, filter_type)

    # 解析读取的笔记内容
    title_list, author_list, \
    keyword_list, publisher_list, \
    publish_time_list, reading_time_list, \
    comment_list, expression_list, \
    importantWord_list, sorted_name_list = readNoteFiles(files, names)

    # 如果不指定日期，默认全部
    if select_flag == 0:
        start_index = 0
        end_index = len(title_list)
        start_date = reading_time_list[0]
        end_date = reading_time_list[-1]
    # 如果指定日期，则获得对应索引范围
    elif select_flag == 1:
        start_index, end_index = getIndexRange(reading_time_list, start_date, end_date)
        if start_index == end_index:
            print("one element", "-", start_index)
            end_index = start_index + 1
        else:
            print("normal range", start_index, "-", end_index)

    # 写入标题
    writeStartHeading(document, 'Notes for Closely Relevant Papers', end_index - start_index, start_date, end_date)
    # 循环写入每个item
    for i in range(start_index, end_index):
        writeItemStyle2(document,
                        title_list[i],
                        author_list[i],
                        keyword_list[i],
                        publisher_list[i],
                        publish_time_list[i],
                        reading_time_list[i],
                        comment_list[i],
                        expression_list[i],
                        importantWord_list[i])
    # 对于expression和word，额外输出到writing文件夹
    outputExpressionAndWord(reading_time_list[start_index:end_index],
                            sorted_name_list[start_index:end_index],
                            expression_list[start_index:end_index],
                            importantWord_list[start_index:end_index])
    # 下面内容另起一页
    document.add_page_break()
    # ----------------------------------------------------------------------------

    # ----------------------------------------------------------------------------
    # Generally Relevant Papers
    # 搜索文件
    paths2, names2, files2 = findAllFiles(search_dir2, filter_type)

    # 读取文件
    title_list2, author_list2, \
    keyword_list2, publisher_list2, \
    publish_time_list2, reading_time_list2, \
    comment_list2, expression_list2, \
    importantWord_list2, sorted_name_list2 = readNoteFiles(files2, names2)

    # 如果不指定日期，默认全部
    if select_flag == 0:
        start_index = 0
        end_index = len(title_list2)
        start_date = reading_time_list2[0]
        end_date = reading_time_list2[-1]
    # 如果指定日期，则获得对应索引范围
    elif select_flag == 1:
        start_index, end_index = getIndexRange(reading_time_list2, start_date, end_date)
        if start_index == end_index:
            print("one element", "-", start_index)
            end_index = start_index + 1
        else:
            print("normal range", start_index, "-", end_index)

    # 写入标题
    writeStartHeading(document, 'Notes for Generally Relevant Papers', end_index - start_index, start_date, end_date)
    # 循环写入每个item
    for i in range(start_index, end_index):
        writeItemStyle2(document,
                        title_list2[i],
                        author_list2[i],
                        keyword_list2[i],
                        publisher_list2[i],
                        publish_time_list2[i],
                        reading_time_list2[i],
                        comment_list2[i],
                        expression_list2[i],
                        importantWord_list2[i])
    # 对于expression和word，额外输出到writing文件夹
    outputExpressionAndWord(reading_time_list2[start_index:end_index],
                            sorted_name_list2[start_index:end_index],
                            expression_list2[start_index:end_index],
                            importantWord_list2[start_index:end_index])
    # 下面内容另起一页
    document.add_page_break()
    # ----------------------------------------------------------------------------

    # ----------------------------------------------------------------------------
    # 统计相关的内容输出
    writeHeading(document, "Statistics for Reading Notes")

    # 获取总共的数量
    total_number = len(reading_time_list) + len(reading_time_list2)
    document.add_paragraph(total_number.__str__() + " papers are read (" +
                           len(reading_time_list).__str__() + " closely relevant papers and " +
                           len(reading_time_list2).__str__() + " generally relevant papers).")

    # 获取第一篇和最后一篇的阅读日期
    start_year, start_month, start_day = parseDateToNumber(reading_time_list[0])
    if reading_time_list[-1] > reading_time_list2[-1]:
        end_year, end_month, end_day = parseDateToNumber(reading_time_list[-1])
    else:
        end_year, end_month, end_day = parseDateToNumber(reading_time_list2[-1])

    d1 = datetime.datetime(start_year, start_month, start_day)
    d2 = datetime.datetime(end_year, end_month, end_day)
    interval = d2 - d1

    # 计算平均的统计信息
    speed_day = 1.0 * total_number / interval.days
    speed_week = speed_day * 7
    speed_month = speed_day * 30

    # 写到文档里
    document.add_paragraph('\tMean Daily Read: ' + round(speed_day, 3).__str__() + " papers" +
                           '\n\tWeekly: ' + round(speed_week, 3).__str__() + " papers" +
                           '\n\tMonthly: ' + round(speed_month, 3).__str__() + " papers")


    joined_dates = []
    joined_dates.extend(reading_time_list)
    joined_dates.extend(reading_time_list2)
    joined_dates.sort()
    start_year, start_month, _ = parseDateToNumber(joined_dates[0])
    end_year, end_month, _ = parseDateToNumber(joined_dates[-1])
    end_day = calendar.monthrange(end_year, end_month)[1]
    joined_timestamps = []

    # 根据起始和终止时间，按月构造时间间隔
    num = (end_year - start_year) * 12 + (end_month - start_month) + 1
    monthly_range = []
    for i in range(num):
        tmp_month = start_month + i
        if tmp_month % 12 == 0:
            d_year = int(tmp_month / 12) - 1
            d_month = 12
        else:
            d_year = int(tmp_month / 12)
            d_month = tmp_month - 12 * d_year
        cur_year = start_year + d_year
        cur_month = d_month
        cur_month_start = 1
        cur_month_end = calendar.monthrange(cur_year, cur_month)[1]
        monthly_range.append([[cur_year, cur_month, cur_month_start], [cur_year, cur_month, cur_month_end]])

    # 将日期转换为时间戳，以方便后续比较
    for i in range(len(joined_dates)):
        joined_timestamps.append(strCvtToTimeStamp(joined_dates[i]))

    # 统计每个时间段内阅读的数量
    monthly_num = []
    monthly_label = []
    for i in range(len(monthly_range)):
        cur_year, cur_month, _ = monthly_range[i][0]
        str_label = str(cur_year) + '.' + str(cur_month).zfill(2)

        start_time = numCvtToTimeStamp2(monthly_range[i][0])
        end_time = numCvtToTimeStamp2(monthly_range[i][1])

        start_index = 0
        end_index = len(joined_timestamps) - 1

        if end_time < joined_timestamps[0]:
            num = 0
        elif start_time > joined_timestamps[-1]:
            num = 0
        else:
            for j in range(len(joined_timestamps)):
                if joined_timestamps[j] >= start_time:
                    start_index = j
                    break
            for k in range(len(joined_timestamps) - 1, -1, -1):
                if joined_timestamps[k] <= end_time:
                    end_index = k + 1
                    break
            num = abs(end_index - start_index)
        monthly_num.append(num)
        monthly_label.append(str_label)

    # 绘制柱状图并保存
    plt.figure(figsize=(12, 8))
    plt.plot(monthly_num, color='orange')
    plt.bar(monthly_label, monthly_num)
    plt.xticks(rotation=90)
    plt.tight_layout()

    t = time.strftime('%Y.%m.%d', time.localtime(time.time()))
    if len(sys.argv) == 1 or len(sys.argv) == 3:
        # 如果选择了某个时间段，采用下面的名字
        if select_flag:
            outpath = cur_dir + "\\Others\\Paper Notes(" + start_date + " - " + end_date + ") - " + t + ".docx"
        else:
            outpath = cur_dir + "\\Others\\PaperNotes(All) - " + t + ".png"
    else:
        if select_flag:
            outpath = sys.argv[1] + "\\Others\\Paper Notes(" + start_date + " - " + end_date + ") - " + t + ".docx"
        else:
            outpath = sys.argv[1] + "\\Others\\PaperNotes(All) - " + t + ".png"
    plt.savefig(outpath, dpi=600)

    # 将绘制的图片插入文档中
    picture = document.add_picture(outpath)
    picture.height = Cm(10)
    picture.width = Cm(15)
    paragraph = document.add_paragraph('Plot of monthly read papers.')
    paragraph.alignment = 1
    # ----------------------------------------------------------------------------

    # 获取当前日期并保存文件
    t = time.strftime('%Y.%m.%d', time.localtime(time.time()))
    if len(sys.argv) == 1 or len(sys.argv) == 3:
        # 如果选择了某个时间段，采用下面的名字
        if select_flag:
            outpath = cur_dir + "\\Notes\\Paper Notes(" + start_date + " - " + end_date + ") - " + t + ".docx"
        else:
            outpath = cur_dir + "\\Notes\\Paper Notes(All) - " + t + ".docx"
    else:
        if select_flag:
            outpath = sys.argv[1] + "\\Notes\\Paper Notes(" + start_date + " - " + end_date + ") - " + t + ".docx"
        else:
            outpath = sys.argv[1] + "\\Notes\\Paper Notes(All) - " + t + ".docx"
    document.save(outpath)

    # 调用Win32 API更新Word文件目录
    word = win32com.client.DispatchEx("Word.Application")
    # 注意传给Word的文件路径必须是绝对路径
    doc = word.Documents.Open(outpath)
    doc.TablesOfContents(1).Update()
    doc.Close(SaveChanges=True)
    word.Quit()
